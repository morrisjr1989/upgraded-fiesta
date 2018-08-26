from docx import Document
from docx.shared import Inches

document = Document()
document.add_paragraph('Dear Hiring Manager:')

p = document.add_paragraph('Thanks for taking the time /t
to look into my resume")

document.save('test.docx')